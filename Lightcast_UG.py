import streamlit as st
import requests
import json
# from dotenv import load_dotenv
import PyPDF2
import os
import re
from supabase import create_client, Client
from openai import OpenAI
import json
from docx import Document
import pandas as pd
from io import BytesIO
import datetime
import zipfile
import toml

def check_lightcast_credentials(client_id, client_secret):
    auth_url = "https://auth.emsicloud.com/connect/token"
    payload = {
        'client_id': client_id,
        'client_secret': client_secret,
        'grant_type': 'client_credentials',
        'scope': 'emsi_open'
    }
    headers = {'Content-Type': 'application/x-www-form-urlencoded'}
    auth_response = requests.post(auth_url, data=payload, headers=headers)

    if auth_response.status_code == 200:
        return True, auth_response.json()['access_token']
    else:
        return False, auth_response.text


def read_pdf(file):
    """
    Function to read a PDF file and return its text content.
    """
    full_text = []
    reader = PyPDF2.PdfReader(file)  # Directly pass the file-like object
    for page_num in range(len(reader.pages)):
        page = reader.pages[page_num]
        full_text.append(page.extract_text())
    return '\n'.join(full_text)


def read_docx(file_path):
    """
    Function to read a .docx file and return its text content, including text within tables.
    """
    doc = Document(file_path)
    full_text = []

    # Extract text from paragraphs
    for para in doc.paragraphs:
        full_text.append(para.text)

    # Extract text from tables
    try:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
    except Exception as e:
        pass
    return '\n'.join(full_text)


def extract_course_details(client, document_text):
    question1 = f"""
    Extract the following information from the course syllabus in the below JSON format:
    "CPS_UG_Course_Code": 
    "CPS_UG_Course_Title": 
    "CPS_UG_Course_Instructor_Name": 
    "CPS_UG_Course_Instructor_Email": 
    "CPS_UG_Course_Term":
    "CPS_UG_Course_Length": 
    "CPS_UG_Course_Credit_Hours": 
    "CPS_UG_Course_Format": 
    "CPS_UG_Course_Description": 
    "CPS_UG_Course_Learning_Outcomes": 
    "CPS_UG_Course_Schedule": Make sure to extract the complete detailed schedule without fail even if its tentative.

    The course syllabus is as follows:
    {document_text}
    """

    stream = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[
            {"role": "user", "content": question1},
        ],
        # stream= True,
    )
    # response = st.write_stream(stream)
    pattern = re.compile(r"```json(.*?)```", re.DOTALL)
    matches = pattern.findall(stream.choices[0].message.content)
    combined_code = "".join(matches)
    combined_code = combined_code.strip()
    try:
        data = json.loads(combined_code)
    except:
        stream = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "user", "content": question1},
            ],
            # stream= True,
        )
        # response = st.write_stream(stream)
        pattern = re.compile(r"```json(.*?)```", re.DOTALL)
        matches = pattern.findall(stream.choices[0].message.content)
        combined_code = "".join(matches)
        combined_code = combined_code.strip()
        data = json.loads(combined_code)

    return data


def main():
    # load_dotenv('Lightcast.env')
    st.set_page_config(layout="wide")

    SUPABASE_URL = st.secrets['SUPABASE_URL']
    SUPABASE_KEY = st.secrets['SUPABASE_KEY']
    supabase: Client = create_client(SUPABASE_URL, SUPABASE_KEY)

    client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

    st.title('CPS Skills Extraction Portal')

    # Create two tabs
    tab1, tab2, tab3 = st.tabs(["Syllabi Skill Extractor", "Lightcast Skill Extractor", "Skills Viewer"])

    with tab1:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            # Streamlit UI
            st.subheader('Lightcast Skills Extractor')

            client_id = st.secrets['CLIENT_ID']
            client_secret = st.secrets['CLIENT_SECRET']
            file_type = st.selectbox('File Type', ('Single Upload', 'Batch Upload'))
            if file_type == 'Single Upload':
                uploaded_files_1 = st.file_uploader("Upload File to Extract Skills:", type=['docx', 'pdf'],
                                                    accept_multiple_files=True)
            elif file_type == 'Batch Upload':
                uploaded_files_2 = st.file_uploader("Upload File to Extract Skills:", type=['zip'])
            confidence_threshold = st.slider("Confidence Threshold", 0.0, 1.0, 0.5)
            course_degree = st.radio('Course Degree', ['Undergraduate', 'Graduate'], horizontal=True)
            # folder_path = st.text_input("Folder Path", value="CPS UG Syllabi")
            db_flag = st.toggle("Upload to the Database?")

            if st.button('Extract Skills'):
                # folder_path = 'CPS UG Syllabi'
                is_valid, response = check_lightcast_credentials(client_id, client_secret)

                if is_valid == False:
                    st.error('Invalid Lightcast Credentials')
                    st.stop()

                if client_id and client_secret:
                    if file_type == 'Single Upload':
                        if uploaded_files_1:
                            for file in uploaded_files_1:
                                # st.write(file)
                                file_name = file.name  # Capture the file name here
                                file_extension = os.path.splitext(file_name)[1].lower()

                                if file_extension == '.pdf':
                                    document_text = read_pdf(file)
                                elif file_extension == '.docx':
                                    document_text = read_docx(file)
                                else:
                                    st.warning('We only support PDF and DOCX files')
                                    st.stop()

                                st.subheader(f"File Name: {file_name}")

                                # Remove text between 'Program Learning Outcomes' and 'Course Learning Outcomes'
                                # modified_text = re.sub(r'Program Learning Outcomes.*?Course Learning Outcomes', 'Course Learning Outcomes', document_text, flags=re.DOTALL)
                                course_info = extract_course_details(client, document_text)
                                st.write('**Course Code:** ', course_info['CPS_UG_Course_Code'])
                                st.write('**Course Title:** ', course_info['CPS_UG_Course_Title'])
                                st.write('**Course Instructor Name:** ', course_info['CPS_UG_Course_Instructor_Name'])
                                st.write('**Course Instructor Email:** ', course_info['CPS_UG_Course_Instructor_Email'])
                                st.write('**Course Degree:** ', course_degree)
                                st.write('**Course Term:** ', course_info['CPS_UG_Course_Term'])
                                st.write('**Course Length:** ', course_info['CPS_UG_Course_Length'])
                                st.write('**Course Credit Hours:** ', course_info['CPS_UG_Course_Credit_Hours'])
                                st.write('**Course Format:** ', course_info['CPS_UG_Course_Format'])
                                st.write('**Course Description:** ', course_info['CPS_UG_Course_Description'])
                                st.write('**Course Learning Outcomes:** ', course_info['CPS_UG_Course_Learning_Outcomes'])
                                st.write('**Course Schedule:** ', course_info['CPS_UG_Course_Schedule'])

                                modified_text = f"Course Description:\n {course_info['CPS_UG_Course_Description']} \n\n Course Learning Outcome:\n {course_info['CPS_UG_Course_Learning_Outcomes']} \n\n Course Schedule:\n {course_info['CPS_UG_Course_Schedule']}"

                                auth_url = "https://auth.emsicloud.com/connect/token"
                                payload = {
                                    'client_id': client_id,
                                    'client_secret': client_secret,
                                    'grant_type': 'client_credentials',
                                    'scope': 'emsi_open'
                                }
                                headers = {'Content-Type': 'application/x-www-form-urlencoded'}
                                auth_response = requests.post(auth_url, data=payload, headers=headers)

                                if auth_response.status_code == 200:
                                    access_token = auth_response.json()['access_token']
                                    skills_url = "https://emsiservices.com/skills/versions/latest/extract/trace"
                                    data_payload = {
                                        "text": modified_text,
                                        "confidenceThreshold": confidence_threshold
                                    }
                                    headers = {
                                        'Authorization': f"Bearer {access_token}",
                                        'Content-Type': "application/json"
                                    }
                                    skills_response = requests.post(skills_url, data=json.dumps(data_payload),
                                                                    headers=headers)

                                    if skills_response.status_code == 200:
                                        skills_response = skills_response.json()
                                        skills = skills_response['data']['skills']
                                        skill_names = [skill['skill']['name'] for skill in skills]
                                        st.write("##### Skills Extracted:\n")
                                        st.write(skill_names)
                                        st.success('Skills successfully extracted.')
                                    else:
                                        st.error(
                                            f"Failed to extract skills. Status code: {skills_response.status_code}")
                                        st.write(skills_response.text)
                                else:
                                    st.error(f"Failed to retrieve token. Status code: {auth_response.status_code}")
                                    st.write(auth_response.text)

                                if db_flag:
                                    # st.write(course_degree)
                                    cps_ug_course_syllabi_data_to_insert = {
                                        "CPS_UG_Course_Code": course_info.get("CPS_UG_Course_Code"),
                                        "CPS_UG_Course_Title": course_info.get("CPS_UG_Course_Title"),
                                        "CPS_UG_Course_Instructor_Name": course_info.get("CPS_UG_Course_Instructor_Name"),
                                        "CPS_UG_Course_Instructor_Email": course_info.get("CPS_UG_Course_Instructor_Email"),
                                        "CPS_UG_Course_Term": course_info.get("CPS_UG_Course_Term"),
                                        "CPS_UG_Course_Length": course_info.get("CPS_UG_Course_Length"),
                                        "CPS_UG_Course_Credit_Hours": course_info.get("CPS_UG_Course_Credit_Hours"),
                                        "CPS_UG_Course_Format": course_info.get("CPS_UG_Course_Format"),
                                        "CPS_UG_Course_Description": course_info.get("CPS_UG_Course_Description"),
                                        "CPS_UG_Course_Learning_Outcomes": course_info.get("CPS_UG_Course_Learning_Outcomes"),
                                        "CPS_UG_Course_Schedule": course_info.get("CPS_UG_Course_Schedule"),
                                        "CPS_UG_Course_Syllabi": document_text.replace('\u0000', ''),
                                        "CPS_UG_Course_Skills": skill_names,
                                        "CPS_UG_Course_File_Name": file_name,
                                        "CPS_UG_Course_Degree": course_degree,
                                    }

                                    # Insert data into CPS_Course_Syllabi
                                    cl_response = supabase.table("CPS_Course_Syllabi_Skills").insert(cps_ug_course_syllabi_data_to_insert).execute()
                                    st.success('Uploaded Successfully!')

                    elif file_type == 'Batch Upload':
                        if uploaded_files_2:
                            with zipfile.ZipFile(uploaded_files_2, 'r') as zip_ref:
                                for file_info in zip_ref.infolist():
                                    # Check if the file is not a directory
                                    if not file_info.is_dir():
                                        file_extension = os.path.splitext(file_info.filename)[1].lower()
                                        if file_info.filename.startswith('~$'):
                                            continue

                                        if file_extension in [".docx", ".pdf"]:

                                            with zip_ref.open(file_info.filename) as file:
                                                file_content = file.read()
                                                file_like_object = BytesIO(file_content)

                                                if file_extension == ".docx":
                                                    document_text = read_docx(file_like_object)

                                                elif file_extension == ".pdf":
                                                    document_text = read_pdf(file_like_object)

                                                st.subheader(f"File Name: {file_info.filename}")

                                                # Remove text between 'Program Learning Outcomes' and 'Course Learning Outcomes'
                                                # modified_text = re.sub(r'Program Learning Outcomes.*?Course Learning Outcomes', 'Course Learning Outcomes', document_text, flags=re.DOTALL)
                                                course_info = extract_course_details(client, document_text)
                                                st.write('**Course Code:** ', course_info['CPS_UG_Course_Code'])
                                                st.write('**Course Title:** ', course_info['CPS_UG_Course_Title'])
                                                st.write('**Course Instructor Name:** ', course_info['CPS_UG_Course_Instructor_Name'])
                                                st.write('**Course Instructor Email:** ', course_info['CPS_UG_Course_Instructor_Email'])
                                                st.write('**Course Degree:** ', course_degree)
                                                st.write('**Course Term:** ', course_info['CPS_UG_Course_Term'])
                                                st.write('**Course Length:** ', course_info['CPS_UG_Course_Length'])
                                                st.write('**Course Credit Hours:** ', course_info['CPS_UG_Course_Credit_Hours'])
                                                st.write('**Course Format:** ', course_info['CPS_UG_Course_Format'])
                                                st.write('**Course Description:** ', course_info['CPS_UG_Course_Description'])
                                                st.write('**Course Learning Outcomes:** ', course_info['CPS_UG_Course_Learning_Outcomes'])
                                                st.write('**Course Schedule:** ', course_info['CPS_UG_Course_Schedule'])

                                                modified_text = f"Course Description:\n {course_info['CPS_UG_Course_Description']} \n\n Course Learning Outcome:\n {course_info['CPS_UG_Course_Learning_Outcomes']} \n\n Course Schedule:\n {course_info['CPS_UG_Course_Schedule']}"

                                                auth_url = "https://auth.emsicloud.com/connect/token"
                                                payload = {
                                                    'client_id': client_id,
                                                    'client_secret': client_secret,
                                                    'grant_type': 'client_credentials',
                                                    'scope': 'emsi_open'
                                                }
                                                headers = {'Content-Type': 'application/x-www-form-urlencoded'}
                                                auth_response = requests.post(auth_url, data=payload, headers=headers)

                                                if auth_response.status_code == 200:
                                                    access_token = auth_response.json()['access_token']
                                                    skills_url = "https://emsiservices.com/skills/versions/latest/extract/trace"
                                                    data_payload = {
                                                        "text": modified_text,
                                                        "confidenceThreshold": confidence_threshold
                                                    }
                                                    headers = {
                                                        'Authorization': f"Bearer {access_token}",
                                                        'Content-Type': "application/json"
                                                    }
                                                    skills_response = requests.post(skills_url, data=json.dumps(data_payload), headers=headers)

                                                    if skills_response.status_code == 200:
                                                        skills_response = skills_response.json()
                                                        skills = skills_response['data']['skills']
                                                        skill_names = [skill['skill']['name'] for skill in skills]
                                                        st.write("##### Skills Extracted:\n")
                                                        st.write(skill_names)
                                                        st.success('Skills successfully extracted.')
                                                    else:
                                                        st.error(
                                                            f"Failed to extract skills. Status code: {skills_response.status_code}")
                                                        st.write(skills_response.text)
                                                else:
                                                    st.error(
                                                        f"Failed to retrieve token. Status code: {auth_response.status_code}")
                                                    st.write(auth_response.text)

                                                if db_flag:
                                                    # st.write(course_degree)

                                                    cps_ug_course_syllabi_data_to_insert = {
                                                        "CPS_UG_Course_Code": course_info.get("CPS_UG_Course_Code"),
                                                        "CPS_UG_Course_Title": course_info.get("CPS_UG_Course_Title"),
                                                        "CPS_UG_Course_Instructor_Name": course_info.get("CPS_UG_Course_Instructor_Name"),
                                                        "CPS_UG_Course_Instructor_Email": course_info.get("CPS_UG_Course_Instructor_Email"),
                                                        "CPS_UG_Course_Term": course_info.get("CPS_UG_Course_Term"),
                                                        "CPS_UG_Course_Length": course_info.get("CPS_UG_Course_Length"),
                                                        "CPS_UG_Course_Credit_Hours": course_info.get("CPS_UG_Course_Credit_Hours"),
                                                        "CPS_UG_Course_Format": course_info.get("CPS_UG_Course_Format"),
                                                        "CPS_UG_Course_Description": course_info.get("CPS_UG_Course_Description"),
                                                        "CPS_UG_Course_Learning_Outcomes": course_info.get("CPS_UG_Course_Learning_Outcomes"),
                                                        "CPS_UG_Course_Schedule": course_info.get("CPS_UG_Course_Schedule"),
                                                        "CPS_UG_Course_Syllabi": document_text.replace('\u0000', ''),
                                                        "CPS_UG_Course_Skills": skill_names,
                                                        "CPS_UG_Course_File_Name": file_info.filename,
                                                        "CPS_UG_Course_Degree": course_degree,
                                                    }
                                                    # Insert data into CPS_Course_Syllabi
                                                    cl_response = supabase.table("CPS_Course_Syllabi_Skills").insert(
                                                        cps_ug_course_syllabi_data_to_insert).execute()
                                                    st.success('Uploaded Successfully!')
                                        else:
                                            st.write(f"{file_info.filename} is not a DOCX or PDF file. Skipping...")

                else:
                    st.warning("Please enter all required fields.")

    with tab2:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            uploaded_files = st.file_uploader("Upload File to Extract Skills:", type=['docx', 'pdf'], accept_multiple_files=True, key='Extract Lightcast Skills')
            if st.button("Process"):
                with st.spinner("Processing"):
                    if uploaded_files:
                        for file in uploaded_files:
                            # st.write(file)
                            file_name = file.name  # Capture the file name here
                            file_extension = os.path.splitext(file_name)[1].lower()

                            if file_extension == '.pdf':
                                document_text = read_pdf(file)
                            elif file_extension == '.docx':
                                document_text = read_docx(file)
                            else:
                                st.warning('We only support PDF and DOCX files')
                                st.stop()

                            auth_url = "https://auth.emsicloud.com/connect/token"
                            payload = {
                                'client_id': client_id,
                                'client_secret': client_secret,
                                'grant_type': 'client_credentials',
                                'scope': 'emsi_open'
                            }
                            headers = {'Content-Type': 'application/x-www-form-urlencoded'}
                            auth_response = requests.post(auth_url, data=payload, headers=headers)

                            if auth_response.status_code == 200:
                                access_token = auth_response.json()['access_token']
                                skills_url = "https://emsiservices.com/skills/versions/latest/extract/trace"
                                data_payload = {
                                    "text": document_text,
                                    "confidenceThreshold": confidence_threshold
                                }
                                headers = {
                                    'Authorization': f"Bearer {access_token}",
                                    'Content-Type': "application/json"
                                }
                                skills_response = requests.post(skills_url, data=json.dumps(data_payload),
                                                                headers=headers)

                                if skills_response.status_code == 200:
                                    skills_response = skills_response.json()
                                    skills = skills_response['data']['skills']
                                    skill_names = [skill['skill']['name'] for skill in skills]
                                    st.write("##### Skills Extracted:\n")
                                    st.write(skill_names)
                                    st.success('Skills successfully extracted.')
                                else:
                                    st.error(f"Failed to extract skills. Status code: {skills_response.status_code}")
                                    st.write(skills_response.text)
                            else:
                                st.error(f"Failed to retrieve token. Status code: {auth_response.status_code}")
                                st.write(auth_response.text)

                            cps_lightcast_skills_data_to_insert = {
                                "CPS_Lightcast_File_Name": file_name,  # Now file_name is properly defined
                                "CPS_Lightcast_File_Text": document_text.replace('\u0000', ''),
                                "CPS_Lightcast_Extracted_Skills": skill_names,
                            }

                            # Insert data into CPS_Course_Syllabi
                            cl_response = supabase.table("CPS_Lightcast_Skills").insert(
                                cps_lightcast_skills_data_to_insert).execute()
                            st.success('Uploaded Successfully!')

    with tab3:
        st.subheader('Skills Viewer')
        cps_table = st.radio('Select Skills Data', ["CPS_Course_Syllabi_Skills", 'CPS_Lightcast_Skills'],
                             horizontal=True)
        start_date = st.date_input("Start date")
        end_date = st.date_input("End date")

        if st.button("Fetch Data"):
            # Convert selected dates to strings in the appropriate format
            start_date_str = start_date.strftime('%Y-%m-%dT%H:%M:%S') + 'Z'
            # end_date_str = end_date.strftime('%Y-%m-%dT%H:%M:%S') + 'Z'

            end_date_plus_one = end_date + datetime.timedelta(days=1)
            end_date_str = end_date_plus_one.strftime('%Y-%m-%dT%H:%M:%S') + 'Z'

            if cps_table == 'CPS_Course_Syllabi_Skills':
                column_names = "CPS_UG_Syllabi_ID, CPS_UG_Course_Code, CPS_UG_Course_Title, CPS_UG_Course_Instructor_Name, CPS_UG_Course_Instructor_Email, CPS_UG_Course_Degree, CPS_UG_Course_Term, CPS_UG_Course_Length, CPS_UG_Course_Credit_Hours, CPS_UG_Course_Format, CPS_UG_Course_Skills, CPS_UG_Course_Created_At, CPS_UG_Course_Description, CPS_UG_Course_Learning_Outcomes, CPS_UG_Course_Schedule, CPS_UG_Course_File_Name"
                response = supabase.table("CPS_Course_Syllabi_Skills").select(column_names).gte(
                    'CPS_UG_Course_Created_At', start_date).lte('CPS_UG_Course_Created_At', end_date_plus_one).order(
                    "CPS_UG_Syllabi_ID").execute()
                data = response.data
                if data:
                    df = pd.DataFrame(data)
                    st.dataframe(data)
                    # Add download button for CSV
                    csv = df.to_csv(index=False)
                    st.download_button(label="Download data as CSV", data=csv, file_name='skills_data.csv',
                                       mime='text/csv')
                else:
                    st.warning("No data found for the date range specified.")
            elif cps_table == 'CPS_Lightcast_Skills':
                column_names = "CPS_Lightcast_ID, CPS_Lightcast_File_Name, CPS_Lightcast_File_Text, CPS_Lightcast_Extracted_Skills, CPS_Lightcast_Created_At"
                response = supabase.table("CPS_Lightcast_Skills").select(column_names).gte(
                    'CPS_Lightcast_Created_At', start_date).lte('CPS_Lightcast_Created_At', end_date_plus_one).order(
                    "CPS_Lightcast_ID").execute()
                data = response.data
                if data:
                    df = pd.DataFrame(data)
                    st.dataframe(data)
                    # Add download button for CSV
                    csv = df.to_csv(index=False)
                    st.download_button(label="Download data as CSV", data=csv, file_name='skills_data.csv',
                                       mime='text/csv')
                else:
                    st.warning("No data found for the date range specified.")


if __name__ == "__main__":
    main()
